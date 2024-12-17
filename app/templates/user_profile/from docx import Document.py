from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Study Comparison Table', 0)

# Define the table data
table_data = [
    ['Feature', 'Alsmadi & Almarashdeh (2020)', 'Saleh et al. (2022)', 'Soviany & Ionescu (2018)', 
     'Ren et al. (2015)', 'Akshatha et al. (2022)', 'Joseph et al. (2015)', 'Priyankan & Fernando (2021)', 
     'Tharwat et al. (2018)', 'Wan et al. (2021)', 'Diwan et al. (2023)', 'Austen et al. (2016)', 
     'Yang et al. (2020)', 'Allken et al. (2018)', 'Sharif Razavian et al. (2014)', 'Yosinski et al. (2014)', 
     'Rachmatullah & Supriana (2018)', 'Pudaruth et al. (2020)'],
    ['Deep Learning (DL)', '✔', '✔', '✔', '✔', '✔', '✔', '✔', '✔', '✔', '✔', '✔', '✔', '✔', '✔', '✔', '✔'],
    ['Object Detection', '', '', '✔', '✔', '✔', '✔', '✔', '', '✔', '', '', '', '', '✔', '', ''],
    ['YOLO Algorithm', '', '', '✔', '', '', '✔', '✔', '', '', '✔', '', '', '✔', '', '', ''],
    ['Transfer Learning', '', '✔', '', '', '✔', '', '✔', '✔', '✔', '✔', '✔', '', '✔', '✔', '✔', '', '✔'],
    ['Mobile Application', '', '', '', '', '', '', '✔', '', '', '', '✔', '', '', '', '✔', '✔'],
    ['Accuracy Emphasis', '', '✔', '', '✔', '✔', '✔', '✔', '✔', '✔', '', '✔', '', '✔', '✔', '✔', '', ''],
    ['Data Augmentation', '', '', '', '', '', '', '✔', '', '✔', '', '✔', '✔', '', '✔', '', '✔'],
    ['Environmental Factors', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '', '']
]

# Add a table to the document
table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))

# Populate the table
for i, row in enumerate(table_data):
    for j, cell in enumerate(row):
        table.cell(i, j).text = cell

# Save the document
file_path = 'C:\Users\Wawiworld\Documents\School\CSC172_DataMining'
doc.save(file_path)

file_path
